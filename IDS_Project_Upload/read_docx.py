#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取 Word 文档文本内容
"""
import sys
try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)

doc = Document('D:/Users/zxacd/source/毕设/毕业论文初稿.docx')

print("=" * 70)
print("论文标题:", doc.core_properties.title)
print("=" * 70)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i}] {text}")

# 提取表格
print("\n" + "=" * 70)
print("表格内容:")
print("=" * 70)
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
